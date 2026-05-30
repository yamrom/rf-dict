#!/usr/bin/env python3
"""
map_entry_citations.py — Map each SL entry to its cited literary work(s)

For each entry in the SL dictionary, extracts the citation(s) from
the Body Text paragraphs and maps them to bibliography SL_codes.

Output: pipeline/entry_citations.json
  {
    "Б-41": {
      "citations": ["Шолохов 2"],
      "authors":   ["Шолохов"],
      "sl_codes":  ["Шолохов-2"]
    },
    ...
  }

Usage:
  python pipeline/map_entry_citations.py
  python pipeline/map_entry_citations.py --entry Б-41
  python pipeline/map_entry_citations.py --author Шолохов  # all entries citing this author
  python pipeline/map_entry_citations.py --match           # match with FR translations
"""

import re
import csv
import json
import argparse
from pathlib import Path
from docx import Document
from collections import defaultdict

PIPELINE_DIR = Path(__file__).parent
REPO_ROOT    = PIPELINE_DIR.parent
SL_DOCX      = Path('/Users/yamrom/work/rf_dict/SL/docs/Sophia_Lubensky_Dictionary 1.docx')
BIBLIO_CSV   = REPO_ROOT / "bibliography" / "sl_bibliography.csv"
OUTPUT_JSON  = PIPELINE_DIR / "entry_citations.json"
FR_CACHE     = PIPELINE_DIR / "fr_translation_cache.json"

ENTRY_PATTERN = re.compile(r'^([А-ЯЁA-Z])-(\d+)\s*[•·]')
RU_CITE       = re.compile(r'\(([А-ЯЁа-яё][А-ЯЁа-яёA-Za-z\-]+(?:\s+[А-ЯЁа-яё][А-ЯЁа-яёA-Za-z\-]+)*)\s+(\d+[a-z]?)\)')


def load_bibliography() -> dict:
    """Load bibliography as dict: surname → list of SL_codes."""
    surname_to_codes = defaultdict(list)
    code_to_row      = {}

    with open(BIBLIO_CSV, encoding='utf-8') as f:
        for row in csv.DictReader(f):
            code   = row['SL_code'].strip()
            author = row.get('Author_RU', '').strip()
            if author:
                surname = author.split()[0]
                surname_to_codes[surname].append(code)
            code_to_row[code] = row

    return surname_to_codes, code_to_row


def build_entry_index(doc) -> dict:
    """Build entry_id → paragraph_index mapping."""
    index = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Normal':
            m = ENTRY_PATTERN.match(p.text.strip())
            if m:
                eid = f"{m.group(1)}-{m.group(2)}"
                if eid not in index:
                    index[eid] = i
    return index


def extract_citations_for_entry(doc, entry_id: str, index: dict) -> list:
    """Extract all citations from Body Text paragraphs of one entry."""
    if entry_id not in index:
        return []

    start = index[entry_id]
    positions = sorted(index.values())
    pos = positions.index(start)
    end = positions[pos+1] if pos+1 < len(positions) else len(doc.paragraphs)

    citations = []
    for p in doc.paragraphs[start:end]:
        if p.style.name == 'Body Text' and p.text.strip():
            for m in RU_CITE.finditer(p.text):
                author_part = m.group(1).strip()
                work_num    = m.group(2).strip()
                cite_str    = f"{author_part} {work_num}"
                if cite_str not in citations:
                    citations.append(cite_str)

    return citations


def match_citation_to_bibliography(citation: str,
                                   surname_to_codes: dict) -> list:
    """
    Match a citation string like 'Шолохов 2' to bibliography SL_codes.
    Returns list of matching SL_codes.
    """
    parts = citation.split()
    if not parts:
        return []

    # Try matching by surname (first word of citation)
    surname = parts[0]
    codes   = surname_to_codes.get(surname, [])

    if len(parts) >= 2:
        # Try to match work number too
        # SL_code format: "Шолохов-2" — number after dash
        work_num = parts[-1].rstrip('abcdefghijklmnopqrstuvwxyz')
        specific = [c for c in codes
                    if c.endswith(f'-{work_num}')]
        if specific:
            return specific

    return codes


def build_full_map(doc, index: dict,
                   surname_to_codes: dict) -> dict:
    """Build complete entry → citations mapping."""
    result = {}

    sorted_entries = sorted(index.keys(),
                            key=lambda e: (e.split('-')[0],
                                          int(e.split('-')[1])
                                          if e.split('-')[1].isdigit() else 0))

    for entry_id in sorted_entries:
        citations = extract_citations_for_entry(doc, entry_id, index)

        authors  = []
        sl_codes = []
        for cite in citations:
            surname = cite.split()[0]
            if surname not in authors:
                authors.append(surname)
            codes = match_citation_to_bibliography(cite, surname_to_codes)
            for code in codes:
                if code not in sl_codes:
                    sl_codes.append(code)

        result[entry_id] = {
            'citations': citations,
            'authors':   authors,
            'sl_codes':  sl_codes,
        }

    return result


def find_entries_with_fr_translation(entry_map: dict,
                                     fr_cache: dict,
                                     code_to_row: dict) -> list:
    """
    Find entries where the cited work has a French translation available.
    Returns list of dicts with entry_id, citation, fr_title, fr_url etc.
    """
    results = []

    for entry_id, data in entry_map.items():
        for sl_code in data['sl_codes']:
            # Check if this SL_code has a French translation in cache
            fr_data = fr_cache.get(sl_code, {})
            if not fr_data.get('found'):
                continue

            translations = fr_data.get('translations', [])
            if not translations:
                continue

            # Find freely available ones first
            free_trans = [t for t in translations if t.get('freely_available')]
            all_trans  = free_trans or translations

            bib_row = code_to_row.get(sl_code, {})

            results.append({
                'entry_id':    entry_id,
                'sl_code':     sl_code,
                'author_ru':   bib_row.get('Author_RU', ''),
                'work_ru':     bib_row.get('Work_RU', ''),
                'work_en':     bib_row.get('Work_EN', ''),
                'citations':   data['citations'],
                'fr_title':    all_trans[0].get('title_fr', ''),
                'fr_translator': all_trans[0].get('translator', ''),
                'fr_year':     all_trans[0].get('year', ''),
                'fr_free':     bool(free_trans),
                'fr_url':      free_trans[0].get('url', '') if free_trans else '',
                'n_fr_translations': len(translations),
            })

    # Sort by freely available first, then by entry_id
    results.sort(key=lambda x: (not x['fr_free'], x['entry_id']))
    return results


if __name__ == '__main__':
    ap = argparse.ArgumentParser(
        description='Map SL entries to cited literary works')
    ap.add_argument('--build',   action='store_true',
                    help='Build the full entry→citation map (slow, reads docx)')
    ap.add_argument('--entry',   help='Show citations for one entry e.g. Б-41')
    ap.add_argument('--author',  help='Show all entries citing an author')
    ap.add_argument('--match',   action='store_true',
                    help='Match entry citations with FR translation cache')
    ap.add_argument('--stats',   action='store_true',
                    help='Show statistics')
    args = ap.parse_args()

    surname_to_codes, code_to_row = load_bibliography()

    if args.build or not OUTPUT_JSON.exists():
        print(f"Loading {SL_DOCX}...")
        doc   = Document(SL_DOCX)
        index = build_entry_index(doc)
        print(f"  {len(index)} entries indexed")

        print("Building citation map...")
        entry_map = build_full_map(doc, index, surname_to_codes)

        with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
            json.dump(entry_map, f, ensure_ascii=False, indent=2)
        print(f"Saved to {OUTPUT_JSON}")

    else:
        print(f"Loading {OUTPUT_JSON}...")
        with open(OUTPUT_JSON, encoding='utf-8') as f:
            entry_map = json.load(f)

    print(f"  {len(entry_map)} entries in map")

    if args.stats:
        total      = len(entry_map)
        with_cite  = sum(1 for v in entry_map.values() if v['citations'])
        with_bib   = sum(1 for v in entry_map.values() if v['sl_codes'])
        no_cite    = total - with_cite

        print(f"\nTotal entries:              {total}")
        print(f"With citations:             {with_cite} "
              f"({100*with_cite//total}%)")
        print(f"Citations matched to biblio:{with_bib} "
              f"({100*with_bib//total}%)")
        print(f"No citations:               {no_cite} "
              f"({100*no_cite//total}%)")

        # Top cited authors
        from collections import Counter
        author_counts = Counter()
        for v in entry_map.values():
            for a in v['authors']:
                author_counts[a] += 1
        print("\nTop 15 cited authors:")
        for author, count in author_counts.most_common(15):
            print(f"  {author:25s} {count:4d}")

    elif args.entry:
        data = entry_map.get(args.entry)
        if not data:
            print(f"Entry {args.entry} not found")
        else:
            print(f"\n{args.entry}:")
            print(f"  Citations: {data['citations']}")
            print(f"  Authors:   {data['authors']}")
            print(f"  SL codes:  {data['sl_codes']}")
            for code in data['sl_codes']:
                row = code_to_row.get(code, {})
                print(f"  {code}: {row.get('Work_RU','')} "
                      f"({row.get('Author_RU','')})")

    elif args.author:
        matches = [(eid, v) for eid, v in entry_map.items()
                   if args.author in v['authors']]
        print(f"\nEntries citing {args.author}: {len(matches)}")
        for eid, v in matches[:20]:
            print(f"  {eid}: {v['citations']}")

    elif args.match:
        if not FR_CACHE.exists():
            print(f"FR translation cache not found: {FR_CACHE}")
            print("Run find_fr_translations.py --lookup first")
            exit(1)

        with open(FR_CACHE, encoding='utf-8') as f:
            fr_cache = json.load(f)

        print(f"FR cache: {len(fr_cache)} entries looked up")
        matches = find_entries_with_fr_translation(
            entry_map, fr_cache, code_to_row)

        free_matches = [m for m in matches if m['fr_free']]
        all_matches  = matches

        print(f"\nEntries with FR translation available: {len(all_matches)}")
        print(f"  — freely available online:           {len(free_matches)}")
        print(f"\nFreely available (sample):")
        for m in free_matches[:20]:
            print(f"  {m['entry_id']:8s} {m['author_ru']:20s} "
                  f"{m['work_ru'][:30]:30s} → {m['fr_title'][:30]}")
            if m['fr_url']:
                print(f"           {m['fr_url']}")

        # Save match results
        out = PIPELINE_DIR / "entry_fr_matches.json"
        with open(out, 'w', encoding='utf-8') as f:
            json.dump(matches, f, ensure_ascii=False, indent=2)
        print(f"\nFull results saved to {out}")

    else:
        ap.print_help()
        print("\nExamples:")
        print("  python pipeline/map_entry_citations.py --build --stats")
        print("  python pipeline/map_entry_citations.py --entry Б-41")
        print("  python pipeline/map_entry_citations.py --author Шолохов")
        print("  python pipeline/map_entry_citations.py --match")
