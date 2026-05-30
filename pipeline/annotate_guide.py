#!/usr/bin/env python3
"""
annotate_guide.py — Generate annotated Guide with SL entries expanded inline
Output: docs/guide_annotated.docx
"""

import re, sys, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PIPELINE_DIR = Path(__file__).parent
REPO_ROOT    = PIPELINE_DIR.parent
GUIDE_DOCX   = Path('/Users/yamrom/work/rf_dict/SL/docs/Sophia_Lubensky_Guide.docx')
DICT_DOCX    = Path('/Users/yamrom/work/rf_dict/SL/docs/Sophia_Lubensky_Dictionary 1.docx')
OUTPUT_DOCX  = REPO_ROOT / "docs" / "guide_annotated.docx"
INDEX_JSON   = PIPELINE_DIR / "sl_index.json"

ENTRY_ID_PAT = re.compile(r'\b([А-ЯЁA-Z]-\d+)\b')
CAPS_PAT     = re.compile(r'\b([А-ЯЁ][А-ЯЁ́]+(?:[\s/]+[А-ЯЁ][А-ЯЁ́]+){1,5})\b')


def load_idiom_index() -> dict:
    if not INDEX_JSON.exists():
        return {}
    with open(INDEX_JSON, encoding='utf-8') as f:
        return json.load(f).get('idiom_to_entry', {})


def find_entry_for_headword(hw: str, idiom_index: dict) -> str:
    """Look up entry ID for a CAPS headword via the idiom index."""
    hw = hw.strip()
    # Try exact match first
    if hw in idiom_index:
        return idiom_index[hw]['entry']
    # Try lowercase direct lookup (most common case)
    hw_lower = hw.lower()
    if hw_lower in idiom_index:
        return idiom_index[hw_lower]['entry']
    # Try stripping stress marks and retry
    import unicodedata
    hw_plain = ''.join(c for c in unicodedata.normalize('NFD', hw_lower)
                       if unicodedata.category(c) != 'Mn')
    for idiom, val in idiom_index.items():
        idiom_plain = ''.join(c for c in unicodedata.normalize('NFD', idiom.lower())
                              if unicodedata.category(c) != 'Mn')
        if idiom_plain == hw_plain:
            return val['entry']
    return ''


def build_dict_index(dict_doc) -> dict:
    ENTRY_PAT = re.compile(r'^([А-ЯЁA-Z])-(\d+)\s*[•·]')
    index = {}
    for i, p in enumerate(dict_doc.paragraphs):
        if p.style.name == 'Normal':
            m = ENTRY_PAT.match(p.text.strip())
            if m:
                eid = f"{m.group(1)}-{m.group(2)}"
                if eid not in index:
                    index[eid] = i
    return index


def get_entry_paragraphs(dict_doc, entry_id, index):
    if entry_id not in index:
        return []
    start = index[entry_id]
    positions = sorted(index.values())
    pos = positions.index(start)
    end = positions[pos+1] if pos+1 < len(positions) else len(dict_doc.paragraphs)
    return [p for p in dict_doc.paragraphs[start:end] if p.text.strip()]


def add_shading(para, fill="F5F5F5"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_border(para, color="999999"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_run(para, text, bold=False, italic=False, size=10,
            color=None, font="Times New Roman"):
    if not text:
        return
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def insert_entry_box(out_doc, entry_id, dict_doc, index):
    paras = get_entry_paragraphs(dict_doc, entry_id, index)

    # Header
    p = out_doc.add_paragraph()
    add_shading(p, "E0E0E0")
    add_border(p, "666666")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.1)
    add_run(p, f"↓ SL entry {entry_id} ↓",
            bold=True, italic=True, size=9, color=(80,80,80))

    if not paras:
        p2 = out_doc.add_paragraph()
        add_shading(p2, "FFE8E8")
        add_run(p2, f"[Entry {entry_id} not found in dictionary]",
                size=9, color=(180,0,0))
    else:
        for p_src in paras:
            style = p_src.style.name
            text  = p_src.text.strip()
            p2 = out_doc.add_paragraph()
            add_shading(p2, "F5F5F5")
            add_border(p2, "BBBBBB")
            p2.paragraph_format.left_indent  = Inches(0.2)
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after  = Pt(1)
            if style == 'Normal':
                add_run(p2, text, bold=True, size=9)
            elif style == 'Body Text':
                add_run(p2, text, italic=True, size=9, color=(60,60,60))
            elif style == 'Heading 4':
                add_run(p2, text, size=9)
            elif style == 'Heading 6':
                add_run(p2, "< " + text, italic=True, size=9, color=(100,100,100))
            else:
                add_run(p2, text, size=9)

    # Footer
    p3 = out_doc.add_paragraph()
    add_shading(p3, "E0E0E0")
    add_border(p3, "666666")
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(6)
    p3.paragraph_format.left_indent  = Inches(0.1)
    add_run(p3, f"↑ end of SL entry {entry_id} ↑",
            bold=True, italic=True, size=9, color=(80,80,80))


def copy_guide_paragraph(out_doc, p):
    style = p.style.name
    text  = p.text.strip()
    if style == 'Heading 1':
        return out_doc.add_heading(text, level=1)
    elif style == 'Heading 2':
        return out_doc.add_heading(text, level=2)
    elif style == 'Heading 3':
        return out_doc.add_heading(text, level=3)
    elif style == 'List Paragraph':
        np = out_doc.add_paragraph(style='List Bullet')
        np.add_run(text).font.size = Pt(11)
        return np
    else:
        np = out_doc.add_paragraph()
        run = np.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        return np


def normalize_entry_id(eid: str) -> str:
    """Normalize Latin lookalike letters to Cyrillic in entry IDs.
    e.g. Latin 'M-18' → Cyrillic 'М-18', 'B-260' → 'В-260' etc.
    """
    # Map of visually identical Latin → Cyrillic uppercase letters
    LATIN_TO_CYR = {
        'A': 'А', 'B': 'В', 'C': 'С', 'E': 'Е', 'H': 'Н',
        'K': 'К', 'M': 'М', 'O': 'О', 'P': 'Р', 'T': 'Т',
        'X': 'Х', 'Y': 'У',
    }
    if '-' not in eid:
        return eid
    letter, num = eid.split('-', 1)
    letter = LATIN_TO_CYR.get(letter, letter)
    return f"{letter}-{num}"


def extract_entry_ids(text):
    raw = list(set(ENTRY_ID_PAT.findall(text)))
    return [normalize_entry_id(e) for e in raw]


def main():
    print(f"Loading Guide...")
    guide_doc = Document(GUIDE_DOCX)
    print(f"  {len(guide_doc.paragraphs)} paragraphs")

    print(f"Loading Dictionary...")
    dict_doc  = Document(DICT_DOCX)
    print(f"  {len(dict_doc.paragraphs)} paragraphs")
    dict_index = build_dict_index(dict_doc)
    print(f"  {len(dict_index)} entries indexed")

    print(f"Loading idiom index...")
    idiom_index = load_idiom_index()
    print(f"  {len(idiom_index)} idioms in index")

    # Output document
    out_doc = Document()
    section = out_doc.sections[0]
    for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
        setattr(section, attr, int(0.9 * 914400))

    # Title
    t = out_doc.add_heading("GUIDE TO THE DICTIONARY", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = out_doc.add_paragraph()
    add_run(sub, "Annotated edition — SL entries expanded inline",
            italic=True, size=10, color=(100,100,100))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    out_doc.add_paragraph()

    expanded = set()
    n_inserted = 0

    print("\nProcessing Guide...")
    for p in guide_doc.paragraphs:
        text  = p.text.strip()
        style = p.style.name
        if not text:
            continue

        copy_guide_paragraph(out_doc, p)

        # Collect entry IDs to expand
        ids_to_expand = extract_entry_ids(text)

        # Also detect CAPS headword references in body/list paragraphs
        if style in ('Body Text', 'List Paragraph', 'Normal'):
            for caps in CAPS_PAT.findall(text):
                caps = caps.strip()
                if len(caps) < 5:
                    continue
                eid = find_entry_for_headword(caps, idiom_index)
                if eid and eid not in ids_to_expand:
                    ids_to_expand.append(eid)

        # Insert entry boxes (each only once)
        for eid in ids_to_expand:
            if eid not in expanded:
                insert_entry_box(out_doc, eid, dict_doc, dict_index)
                expanded.add(eid)
                n_inserted += 1
                print(f"  Inserted {eid}")

    print(f"\nTotal entries inserted: {n_inserted}")
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    out_doc.save(OUTPUT_DOCX)
    print(f"Saved to {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
