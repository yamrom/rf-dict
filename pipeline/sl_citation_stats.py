#!/usr/bin/env python3
"""
sl_citation_stats.py — Statistics on citation coverage in SL dictionary

Counts:
1. Entries with examples attributed to bibliography authors
2. Entries with examples but no author attribution (авторский пример)
3. Entries with no examples at all

Usage:
  python3 sl_citation_stats.py
  python3 sl_citation_stats.py --details   # show sample entries for each category
  python3 sl_citation_stats.py --letter Б  # restrict to one letter
"""

import re
import csv
import argparse
from pathlib import Path
from collections import Counter, defaultdict
from docx import Document

SL_DOCX  = Path('/Users/yamrom/work/rf_dict/SL/docs/Sophia_Lubensky_Dictionary 1.docx')
BIBLIO   = Path('/Users/yamrom/work/rf_dict/rf_dict/bibliography/sl_bibliography.csv')

ENTRY_PATTERN = re.compile(r'^([А-ЯЁA-Z])-(\d+)\s*[•·]')
RU_CITE       = re.compile(r'\(([А-ЯЁа-яё][А-ЯЁа-яёA-Za-z\-]+\s+\d+[a-z]?)\)')


def load_bibliography():
    """Load author surnames from bibliography CSV."""
    authors = set()
    try:
        with open(BIBLIO, encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Column is Author_RU e.g. "Абрамов Федор Александрович"
                author_ru = row.get('Author_RU', '').strip()
                if author_ru:
                    surname = author_ru.split()[0]
                    authors.add(surname)
    except Exception as e:
        print(f"Warning: could not load bibliography: {e}")
    return authors


def get_entry_paragraphs(doc, index):
    """Yield (entry_id, paragraphs) for each entry."""
    sorted_positions = sorted(index.items(), key=lambda x: x[1])

    for i, (entry_id, start) in enumerate(sorted_positions):
        end = sorted_positions[i+1][1] if i+1 < len(sorted_positions) else len(doc.paragraphs)
        paras = [p for p in doc.paragraphs[start:end] if p.text.strip()]
        yield entry_id, paras


def analyze_entry(entry_id, paras, bib_authors):
    """
    Analyze one entry for citation coverage.
    Returns dict with:
      has_example: bool
      has_attributed: bool
      has_unattributed: bool
      citations: list of citation strings found
      has_ellipsis: bool  — example contains ... indicating cut context
      ellipsis_count: int — number of ellipses in examples
    """
    citations      = []
    has_body       = False
    has_ru_cite    = False
    ellipsis_count = 0
    full_example_text = ""

    for p in paras:
        if p.style.name == 'Body Text':
            text = p.text.strip()
            if not text:
                continue
            has_body = True
            full_example_text += text + " "
            for m in RU_CITE.finditer(text):
                cite = m.group(1)
                citations.append(cite)
                has_ru_cite = True

    # Count ellipses in example text
    # Both "..." and "…" (U+2026) and "…." patterns
    ellipsis_count = (
        full_example_text.count('…') +
        full_example_text.count('...') +
        full_example_text.count('. . .')
    )

    # Check bibliography match
    has_attributed = False
    matched_authors = []
    for cite in citations:
        parts = cite.split()
        if parts:
            surname = parts[0]
            if surname in bib_authors or not bib_authors:
                has_attributed = True
                matched_authors.append(surname)

    return {
        'has_example':      has_body,
        'has_attributed':   has_attributed,
        'has_unattributed': has_body and not has_ru_cite,
        'citations':        citations,
        'matched_authors':  matched_authors,
        'has_ellipsis':     ellipsis_count > 0,
        'ellipsis_count':   ellipsis_count,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--details', action='store_true',
                    help='Show sample entries for each category')
    ap.add_argument('--letter', help='Restrict to one letter e.g. Б')
    ap.add_argument('--top_authors', action='store_true',
                    help='Show most cited authors')
    args = ap.parse_args()

    print(f"Loading {SL_DOCX}...")
    doc = Document(SL_DOCX)
    print(f"  {len(doc.paragraphs)} paragraphs")

    # Build entry index
    index = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Normal':
            m = ENTRY_PATTERN.match(p.text.strip())
            if m:
                eid = f"{m.group(1)}-{m.group(2)}"
                if eid not in index:
                    index[eid] = i
    print(f"  {len(index)} entries indexed")

    # Load bibliography
    bib_authors = load_bibliography()
    print(f"  {len(bib_authors)} bibliography authors loaded")
    print()

    # Analyze all entries
    stats = {
        'total':             0,
        'with_attributed':   0,
        'with_unattributed': 0,
        'no_example':        0,
        'with_example':      0,
        'with_ellipsis':     0,   # attributed examples containing ...
        'without_ellipsis':  0,   # attributed examples without ...
    }

    author_counts = Counter()
    samples = defaultdict(list)

    for entry_id, paras in get_entry_paragraphs(doc, index):
        # Filter by letter if requested
        if args.letter and not entry_id.startswith(args.letter + '-'):
            continue

        result = analyze_entry(entry_id, paras, bib_authors)
        stats['total'] += 1

        if not result['has_example']:
            stats['no_example'] += 1
            if len(samples['no_example']) < 5:
                samples['no_example'].append(entry_id)
        else:
            stats['with_example'] += 1
            if result['has_attributed']:
                stats['with_attributed'] += 1
                if result['has_ellipsis']:
                    stats['with_ellipsis'] += 1
                    if len(samples['with_ellipsis']) < 5:
                        samples['with_ellipsis'].append(
                            f"{entry_id} ({result['citations'][0]}, "
                            f"{result['ellipsis_count']} ellipses)")
                else:
                    stats['without_ellipsis'] += 1
                    if len(samples['without_ellipsis']) < 5:
                        samples['without_ellipsis'].append(
                            f"{entry_id} ({result['citations'][0]})")
                if len(samples['attributed']) < 5:
                    samples['attributed'].append(
                        f"{entry_id} ({result['citations'][0]})")
            else:
                stats['with_unattributed'] += 1
                if len(samples['unattributed']) < 5:
                    samples['unattributed'].append(entry_id)

        for author in result['matched_authors']:
            author_counts[author] += 1

    # Print results
    total = stats['total']
    print(f"{'='*55}")
    print(f"SL CITATION COVERAGE STATISTICS")
    if args.letter:
        print(f"Letter: {args.letter}")
    print(f"{'='*55}")
    print(f"Total entries analyzed:          {total:6d}")
    print()
    print(f"With attributed examples:        {stats['with_attributed']:6d}  "
          f"({100*stats['with_attributed']//total if total else 0}%)")
    print(f"  — with ellipsis (...):         {stats['with_ellipsis']:6d}  "
          f"({100*stats['with_ellipsis']//stats['with_attributed'] if stats['with_attributed'] else 0}% of attributed)")
    print(f"  — without ellipsis:            {stats['without_ellipsis']:6d}  "
          f"({100*stats['without_ellipsis']//stats['with_attributed'] if stats['with_attributed'] else 0}% of attributed)")
    print(f"With unattributed examples:      {stats['with_unattributed']:6d}  "
          f"({100*stats['with_unattributed']//total if total else 0}%)")
    print(f"With any example:                {stats['with_example']:6d}  "
          f"({100*stats['with_example']//total if total else 0}%)")
    print(f"No example at all:               {stats['no_example']:6d}  "
          f"({100*stats['no_example']//total if total else 0}%)")

    if args.details:
        print()
        print("Sample — attributed with ellipsis:")
        for s in samples.get('with_ellipsis', []):
            print(f"  {s}")
        print("Sample — attributed without ellipsis:")
        for s in samples.get('without_ellipsis', []):
            print(f"  {s}")
        print("Sample — unattributed examples:")
        for s in samples['unattributed']:
            print(f"  {s}")
        print("Sample — no example:")
        for s in samples['no_example']:
            print(f"  {s}")

    if args.top_authors:
        print()
        print("Most cited authors (top 20):")
        for author, count in author_counts.most_common(20):
            print(f"  {author:25s} {count:4d}")


if __name__ == '__main__':
    main()
