#!/usr/bin/env python3
"""
sl_docx_parser.py — Parse SL Yale Revised Edition Dictionary docx

Structure confirmed from analysis:
- Entry header: 1-3 consecutive Normal paragraphs (+ Heading 4 continuations)
- Body Text: examples — either full RU+EN in one para, or split across two
- Heading 4: equivalents continuation
- Heading 6: etymological notes
- Empty Body Text: separator between examples

Usage:
  python3 sl_docx_parser.py --entry Б-41
  python3 sl_docx_parser.py --range Б-36 Б-44
  python3 sl_docx_parser.py --stats
  python3 sl_docx_parser.py --entry Б-41 --raw
"""

import re, sys, argparse
from pathlib import Path
from docx import Document

SL_DOCX = Path('/Users/yamrom/work/rf_dict/SL/docs/Sophia_Lubensky_Dictionary 1.docx')

ENTRY_PATTERN = re.compile(r'^([А-ЯЁA-Z])-(\d+)\s*[•·]')
RU_CITE = re.compile(r'\(([А-ЯЁа-яё][А-ЯЁа-яёA-Za-z\-]+\s+\d+[a-z]?)\)')
EN_CITE = re.compile(r'\((\d+[a-z])\)')


def load_doc(path=SL_DOCX):
    return Document(path)


def build_index(doc):
    index = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Normal':
            m = ENTRY_PATTERN.match(p.text.strip())
            if m:
                eid = f"{m.group(1)}-{m.group(2)}"
                if eid not in index:
                    index[eid] = i
    return index


def get_entry_paragraphs(doc, entry_id, index):
    if entry_id not in index:
        return None
    start = index[entry_id]
    positions = sorted(index.values())
    pos = positions.index(start)
    end = positions[pos+1] if pos+1 < len(positions) else len(doc.paragraphs)
    return [{'style': p.style.name, 'text': p.text}
            for p in doc.paragraphs[start:end]]


def rejoin_hyphens(text):
    """Fix hyphenated line-break artifacts: 'say- ing' → 'saying'"""
    return re.sub(r'([А-ЯЁа-яёA-Za-z])-\s+([А-ЯЁа-яёa-z])', r'\1\2', text)


def parse_example(text):
    """Parse Body Text paragraph into RU + EN components."""
    text = text.strip()
    ru_m = RU_CITE.search(text)
    en_m = EN_CITE.search(text)

    if ru_m and en_m and ru_m.start() < en_m.start():
        ru_end = ru_m.end()
        return {
            'ru': text[:ru_end].strip(),
            'en': text[ru_end:].strip(),
            'ru_cite': ru_m.group(1),
            'en_cite': en_m.group(1),
            'status': 'complete'
        }
    elif ru_m:
        return {'ru': text, 'en': '', 'ru_cite': ru_m.group(1), 'en_cite': '', 'status': 'ru_only'}
    elif en_m:
        return {'ru': '', 'en': text, 'ru_cite': '', 'en_cite': en_m.group(1), 'status': 'en_only'}
    else:
        return {'ru': text, 'en': '', 'ru_cite': '', 'en_cite': '', 'status': 'no_cite'}


def parse_entry(entry_id, paragraphs):
    # ── 1. Separate header lines from body ──────────────────────
    header_parts = []
    body_texts   = []
    etym_parts   = []
    in_header    = True

    for para in paragraphs:
        style = para['style']
        text  = para['text'].strip()
        if not text:
            if not in_header:
                body_texts.append(None)   # empty = example separator
            continue

        if style == 'Heading 6':
            etym_parts.append(text)
            in_header = False
        elif in_header and style in ('Normal', 'Heading 4', 'Heading 3'):
            header_parts.append(text)
        elif style == 'Body Text':
            in_header = False
            body_texts.append(text)
        elif style in ('Normal', 'Heading 4') and not in_header:
            # Sense continuation after examples
            body_texts.append(f"[CONT] {text}")
        elif style == 'List Paragraph':
            in_header = False
            body_texts.append(f"[LIST] {text}")

    # ── 2. Build header string ───────────────────────────────────
    full = rejoin_hyphens(' '.join(header_parts))
    full = re.sub(r'^[А-ЯЁA-Z]-\d+\s*[•·]\s*', '', full)

    bracket = full.find('[')
    if bracket > 0:
        head_part = full[:bracket].strip()
        gram_part = full[bracket:]
    else:
        head_part = full
        gram_part = ''

    headwords = [h.strip() for h in re.split(r';\s*', head_part) if h.strip()]

    gm = re.search(r'\[([^\]]+)\]', gram_part)
    grammar_raw = rejoin_hyphens(gm.group(1).strip()) if gm else ''

    # ── 3. Detect phrase type ────────────────────────────────────
    phrase_type = 'unknown'
    for pt in ['saying', 'NP', 'VP', 'AdjP', 'AdvP', 'PrepP',
               'как +', 'Interj', 'formula phrase']:
        if pt.lower() in grammar_raw.lower():
            phrase_type = pt
            break

    # ── 4. Parse senses + equivalents ───────────────────────────
    after_gram = gram_part[gm.end():].strip() if gm else gram_part
    sense_splits = re.split(r'\s+(?=\d+\.\s)', after_gram)

    LABEL_PAT = re.compile(
        r'^((?:coll|iron|lit|obs|obsoles|old-fash|rare|disapprov|'
        r'humor|elev|offic|vulg|rhet|substand|highly\s+coll|'
        r'folk\s+poet|euph|derog|condes|impol|rude|approv|recent|'
        r'mil|dial|special)\s*[,.]?\s*)+', re.IGNORECASE)

    senses = []
    for i, st in enumerate(sense_splits, 1):
        st = st.strip()
        if not st:
            continue
        labels = []
        lm = LABEL_PAT.match(st)
        if lm:
            labels = [l.strip().rstrip('.,')
                      for l in re.split(r'[,\s]+', lm.group(0)) if l.strip()]
            st = st[lm.end():].strip()

        dm = re.search(r'[≃≈=]', st)
        if dm:
            definition  = st[:dm.start()].strip().rstrip(':')
            equiv_text  = st[dm.start()+1:].strip().rstrip('.')
            equivalents = [e.strip()
                           for e in re.split(r';\s*(?=[a-zA-Z(\[])', equiv_text)
                           if e.strip()]
        else:
            definition  = st
            equivalents = []

        senses.append({
            'sense_num':      str(i),
            'usage_labels':   labels,
            'definition_en':  definition,
            'equivalents_en': equivalents,
            'examples':       [],
        })

    if not senses:
        senses = [{'sense_num':'1','usage_labels':[],
                   'definition_en':'','equivalents_en':[],'examples':[]}]

    # ── 5. Parse examples ────────────────────────────────────────
    examples = []
    pending_ru = None
    pending_ru_cite = ''

    for text in body_texts:
        if text is None:          # empty separator
            if pending_ru:
                examples.append({'ru': pending_ru, 'en': '',
                                  'ru_cite': pending_ru_cite, 'en_cite': '',
                                  'status': 'ru_only'})
                pending_ru = None
            continue
        if text.startswith('[CONT]') or text.startswith('[LIST]'):
            continue

        p = parse_example(text)

        if p['status'] == 'complete':
            if pending_ru:
                examples.append({'ru': pending_ru, 'en': '',
                                  'ru_cite': pending_ru_cite, 'en_cite': '',
                                  'status': 'ru_only'})
                pending_ru = None
            examples.append(p)

        elif p['status'] == 'ru_only':
            if pending_ru:
                examples.append({'ru': pending_ru, 'en': '',
                                  'ru_cite': pending_ru_cite, 'en_cite': '',
                                  'status': 'ru_only'})
            pending_ru      = p['ru']
            pending_ru_cite = p['ru_cite']

        elif p['status'] == 'en_only':
            if pending_ru:
                examples.append({'ru': pending_ru, 'en': p['en'],
                                  'ru_cite': pending_ru_cite,
                                  'en_cite': p['en_cite'],
                                  'status': 'split_para'})
                pending_ru = None
            else:
                examples.append(p)   # orphaned EN

        elif p['status'] == 'no_cite':
            if pending_ru:
                pending_ru += ' ' + p['ru']

    if pending_ru:
        examples.append({'ru': pending_ru, 'en': '',
                          'ru_cite': pending_ru_cite, 'en_cite': '',
                          'status': 'ru_only'})

    if examples:
        senses[-1]['examples'] = examples

    return {
        'entry_id':   entry_id,
        'headwords':  headwords,
        'grammar_raw': grammar_raw,
        'phrase_type': phrase_type,
        'senses':      senses,
        'etym_notes':  etym_parts,
    }


def print_entry(entry):
    print(f"\n{'='*65}")
    print(f"Entry:      {entry['entry_id']}")
    print(f"Headwords:  {entry['headwords']}")
    print(f"Grammar:    [{entry['grammar_raw']}]")
    print(f"Type:       {entry['phrase_type']}")
    for s in entry['senses']:
        print(f"\n  Sense {s['sense_num']}:")
        if s['usage_labels']:
            print(f"    Labels:  {s['usage_labels']}")
        if s['definition_en']:
            print(f"    Def:     {s['definition_en'][:80]}")
        for eq in s['equivalents_en'][:5]:
            print(f"    ≃ {eq[:75]}")
        for ex in s['examples']:
            rc = ex.get('ru_cite','')
            ec = ex.get('en_cite','')
            st = ex.get('status','')
            print(f"    [{st}]")
            if ex['ru']:
                print(f"      RU ({rc}): {ex['ru'][:75]}")
            if ex['en']:
                print(f"      EN ({ec}): {ex['en'][:75]}")
    if entry['etym_notes']:
        print(f"\n  Etym: {entry['etym_notes'][0][:80]}")


if __name__ == '__main__':
    ap = argparse.ArgumentParser(description='SL Dictionary docx parser')
    ap.add_argument('--entry', help='Entry ID e.g. Б-41')
    ap.add_argument('--range', nargs=2, metavar=('FROM','TO'))
    ap.add_argument('--stats', action='store_true')
    ap.add_argument('--raw',   action='store_true')
    args = ap.parse_args()

    doc   = load_doc()
    index = build_index(doc)
    print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(index)} entries")

    if args.stats:
        from collections import Counter
        letters = Counter(k.split('-')[0] for k in index)
        print("\nEntries per letter:")
        for l, c in sorted(letters.items()):
            print(f"  {l}: {c}")

    elif args.entry:
        paras = get_entry_paragraphs(doc, args.entry, index)
        if not paras:
            print(f"Not found: {args.entry}"); sys.exit(1)
        if args.raw:
            for p in paras:
                if p['text'].strip():
                    print(f"  [{p['style']:20s}] {p['text'][:95]}")
        else:
            print_entry(parse_entry(args.entry, paras))

    elif args.range:
        from_id, to_id = args.range
        in_range = False
        for eid, _ in sorted(index.items(), key=lambda x: x[1]):
            if eid == from_id: in_range = True
            if in_range:
                print_entry(parse_entry(eid, get_entry_paragraphs(doc, eid, index)))
            if eid == to_id: break

    else:
        ap.print_help()
