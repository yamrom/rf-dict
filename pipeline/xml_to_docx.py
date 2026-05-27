#!/usr/bin/env python3
"""
xml_to_docx.py — Render RF dictionary XML entry to Word document

Usage:
  python3 pipeline/xml_to_docx.py input.xml output.docx
  python3 pipeline/xml_to_docx.py --entry Б-41
"""

import sys
import argparse
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PIPELINE_DIR = Path(__file__).parent
REPO_ROOT    = PIPELINE_DIR.parent
XML_DIR      = REPO_ROOT / "xml" / "entries"
OUT_DIR      = REPO_ROOT / "pdf"   # reuse pdf dir for docx output


def get_text(root, xpath, default=""):
    el = root.find(xpath)
    return (el.text or "").strip() if el is not None else default


def parse_entry_xml(xml_path: str) -> dict:
    """Parse RF entry XML into a flat dict."""
    tree = ET.parse(xml_path)
    root = tree.getroot()

    # French example element
    fr_ex_el = root.find('.//example[@lang="fr"]')

    return {
        "id":            root.get("id", ""),
        "status":        root.get("status", ""),
        "canonical":     get_text(root, ".//canonical"),
        "variants":      [e.text.strip() for e in root.findall(".//variant") if e.text],
        "bracket":       get_text(root, ".//grammar/bracket"),
        "paraphrase_ru": get_text(root, './/grammar/paraphrase[@lang="ru"]'),
        "paraphrase_fr": get_text(root, './/grammar/paraphrase[@lang="fr"]'),
        "stylistic":     get_text(root, "stylistic_labels"),
        "fr_definition": get_text(root, ".//french/definition"),
        "fr_def_type":   (root.find(".//french/definition") or ET.Element("x")).get("type", ""),
        "match_type":    get_text(root, ".//french/match_type"),
        "fr_paraphrase": get_text(root, ".//french/fr_paraphrase"),
        "ru_sentence":   get_text(root, './/example[@lang="ru"]/sentence'),
        "ru_citation":   get_text(root, './/example[@lang="ru"]/citation'),
        "fr_sentence":   get_text(root, './/example[@lang="fr"]/sentence'),
        "fr_status":     fr_ex_el.get("status", "") if fr_ex_el is not None else "",
    }


def add_run(para, text, bold=False, italic=False,
            size=11, color=None, font="Times New Roman"):
    """Add a formatted run to a paragraph."""
    if not text:
        return
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_border_bottom(para):
    """Add a thin bottom border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '2')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def render_entry_to_doc(entry: dict, doc: Document):
    """Add one dictionary entry to the Word document."""

    # ── 1. HEADWORD LINE ────────────────────────────────────
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)

    # Entry number
    add_run(para, f"{entry['id']} ", bold=True, size=11)
    # Bullet
    add_run(para, "• ", bold=False, size=11)
    # Canonical headword
    add_run(para, entry["canonical"], bold=True, size=11)

    # Variants — filter out trivial sub-forms
    if entry["variants"]:
        seen = set()
        seen.add(entry["canonical"])
        for v in entry["variants"]:
            # Skip if too short, already shown, or just a fragment
            if len(v) < 6 or v in seen:
                continue
            # Skip if it's a substring of canonical (likely a fragment)
            if v in entry["canonical"]:
                continue
            seen.add(v)
            add_run(para, "; ", bold=False, size=11)
            add_run(para, v, bold=True, size=11)

    # ── 2. GRAMMAR BRACKET ──────────────────────────────────
    if entry["bracket"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        add_run(para, entry["bracket"], size=10, color=(80, 80, 80))

    # ── 3. STYLISTIC LABEL ──────────────────────────────────
    stylistic = entry.get("stylistic", "")
    if stylistic and stylistic.lower() not in ("", "null", "none"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        add_run(para, stylistic, italic=True, size=10, color=(90, 90, 90))

    # ── 4. FRENCH DEFINITION ────────────────────────────────
    if entry["fr_definition"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)

        if entry["match_type"] == "A":
            # Type A: bold French idiomatic equivalent
            add_run(para, entry["fr_definition"], bold=True, size=11)
            # French paraphrase in smaller italic
            if entry["fr_paraphrase"]:
                add_run(para, " — ", size=10)
                add_run(para, entry["fr_paraphrase"],
                        italic=True, size=10, color=(60, 60, 60))
        else:
            # Type B: regular paraphrase definition
            add_run(para, entry["fr_definition"], size=11)

    # ── 5. PARAPHRASES (review section — smaller, grey) ─────
    if entry["paraphrase_ru"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        add_run(para, "[ru] ", size=9, color=(150, 150, 150))
        add_run(para, entry["paraphrase_ru"],
                italic=True, size=9, color=(120, 120, 120))

    if entry["paraphrase_fr"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        add_run(para, "[fr] ", size=9, color=(150, 150, 150))
        add_run(para, entry["paraphrase_fr"],
                italic=True, size=9, color=(120, 120, 120))

    # ── 6. RUSSIAN EXAMPLE ──────────────────────────────────
    if entry["ru_sentence"]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent   = Inches(0.25)
        para.paragraph_format.space_after   = Pt(1)
        add_run(para, "◆ ", size=11)
        add_run(para, entry["ru_sentence"], italic=True, size=11)
        if entry["ru_citation"]:
            add_run(para, f" ({entry['ru_citation']})",
                    size=10, color=(80, 80, 80))

    # ── 7. FRENCH EXAMPLE ───────────────────────────────────
    if entry["fr_sentence"]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(4)
        add_run(para, entry["fr_sentence"], size=11)
        if entry["fr_status"] == "constructed":
            add_run(para, " [forgé]", italic=True, size=9,
                    color=(0, 112, 192))

    # ── 8. SEPARATOR ────────────────────────────────────────
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    add_border_bottom(para)


def xml_to_docx(xml_path: str, docx_path: str):
    """Full pipeline: XML → Word document."""
    entry = parse_entry_xml(xml_path)

    doc = Document()

    # Page setup — A5, narrow margins (good for dictionary)
    section = doc.sections[0]
    section.page_width  = int(5.83 * 914400)   # A5 width in EMU
    section.page_height = int(8.27 * 914400)   # A5 height in EMU
    section.left_margin   = int(0.75 * 914400)
    section.right_margin  = int(0.75 * 914400)
    section.top_margin    = int(0.75 * 914400)
    section.bottom_margin = int(0.75 * 914400)

    # Default style
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    render_entry_to_doc(entry, doc)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    return docx_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("xml",  nargs="?", help="Input XML file")
    ap.add_argument("docx", nargs="?", help="Output docx file")
    ap.add_argument("--entry", help="Entry ID e.g. Б-41")
    args = ap.parse_args()

    if args.entry:
        letter   = args.entry.split("-")[0]
        xml_path = XML_DIR / letter / f"entry_{args.entry.replace('-','_')}.xml"
        out_path = OUT_DIR / f"entry_{args.entry.replace('-','_')}.docx"
        xml_to_docx(str(xml_path), str(out_path))
    elif args.xml and args.docx:
        xml_to_docx(args.xml, args.docx)
    else:
        ap.print_help()
